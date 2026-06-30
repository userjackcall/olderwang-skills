"""
Helpers for APPENDING brand-new sections (paragraphs and tables) to the end of an
already-generated 探究與實作教案.docx — used for the optional supplementary materials
(WSQ worksheet, 學習表現檢核表, 學習評量+評分規準).

Unlike docx_helpers.py (which edits EXISTING merged-cell tables inside the official
blank template and must avoid python-docx's table.cell() merge-resolution gotcha),
the functions here build entirely NEW paragraphs/tables with no pre-existing merges,
so plain python-docx Document/Table APIs are safe to use throughout.

Usage pattern:
    from docx import Document
    doc = Document('teaching_plan.docx')   # the already-completed lesson plan
    from append_helpers import add_p, page_break, make_bordered_table
    page_break(doc)
    add_p(doc, [('標題文字', True)], align='center')
    ...
    doc.save('teaching_plan.docx')
"""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_ALIGN = {
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
}


def add_run(p, text, bold=False, size=24, italic=False):
    """size is in half-points (24 = 12pt), matching the template's default body text size."""
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size / 2)
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '標楷體')
    r.font.bold = bold
    r.font.italic = italic
    return r


def add_p(doc, text_runs, align=None, space_after=6, space_before=0, indent=None):
    """text_runs: list of (text, bold) tuples, concatenated into one paragraph."""
    p = doc.add_paragraph()
    if align:
        p.alignment = _ALIGN.get(align, align)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    for text, bold in text_runs:
        add_run(p, text, bold=bold)
    return p


def page_break(doc):
    """Insert a page break so each supplementary section starts on its own page."""
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)


def make_bordered_table(doc, n_rows, n_cols, col_widths_dxa, total_width_dxa=9747):
    """Create a new table styled to match the rest of the lesson-plan template:
    single 4pt black borders, fixed layout, centered, explicit column widths (dxa units,
    1440 dxa = 1 inch). col_widths_dxa must have n_cols entries and should sum to
    roughly total_width_dxa for a clean fit on the page."""
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tblPr = table._tbl.find(qn('w:tblPr'))
    tblW = tblPr.find(qn('w:tblW'))
    tblW.set(qn('w:w'), str(total_width_dxa))
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    tblPr.append(tblBorders)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_widths_dxa[i]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
    return table


def style_header_row(table, headers, shade='D9D9D9', font_size=22):
    """Fill + bold + shade a table's first row as a header. Uses docx_helpers.set_tc_text/
    set_tc_shade on the raw tc element (table.rows[0].cells[i]._tc) since this is a brand
    new, non-merged table — safe to mix with the merge-safe helpers."""
    from docx_helpers import set_tc_text, set_tc_shade
    for i, htext in enumerate(headers):
        tc = table.rows[0].cells[i]._tc
        set_tc_text(tc, htext, bold=True, font_size=font_size, align='center')
        set_tc_shade(tc, shade)


def fill_table_rows(table, rows_data, start_row=1, font_size=20, center_cols=None):
    """rows_data: list of tuples, one per data row, len == n_cols.
    center_cols: optional set of column indices to center-align (others left-aligned)."""
    from docx_helpers import set_tc_text
    center_cols = center_cols or set()
    for r_offset, data in enumerate(rows_data):
        ridx = start_row + r_offset
        for cidx, val in enumerate(data):
            tc = table.rows[ridx].cells[cidx]._tc
            align = 'center' if cidx in center_cols else None
            set_tc_text(tc, val, font_size=font_size, align=align)
