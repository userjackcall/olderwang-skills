"""
highlight_table6.py — shade selected 學習表現子項 rows in table6 of a lesson plan docx.

Table6 is pre-printed static text and must NOT have its text edited.
The only permitted operation is shading selected cells yellow (FFE699)
to indicate which 學習表現 items apply to this lesson.

IMPORTANT: This step is REQUIRED and easy to forget because table6 looks
"complete" without any shading (the text is pre-printed). Leaving it unshaded
makes ALL sub-items appear unselected, which is incorrect.
Call highlight_table6_in_doc() AFTER filling all other tables (especially table7
whose 子項 column drives the selection) and BEFORE saving the docx.

Usage:
    from docx import Document
    from docx_helpers import get_tr_list, get_tc, set_tc_shade
    from highlight_table6 import highlight_table6_in_doc

    doc = Document('lesson_plan.docx')
    # ... fill all other tables ...
    selected = {"推理論證", "建立模型", "計畫與執行", "培養科學探究的興趣"}
    highlight_table6_in_doc(doc, selected)
    doc.save('lesson_plan.docx')

Alternatively, derive the selected set automatically from table7's 子項 column:
    selected = collect_selected_from_table7(doc)
    highlight_table6_in_doc(doc, selected)
"""

from docx_helpers import get_tr_list, get_tc, set_tc_shade

HIGHLIGHT = 'FFE699'

# table6 col1 (學習表現子項) text → raw row index (0-indexed <w:tr> list)
SUBITEM_ROW = {
    '想像創造': 1,
    '推理論證': 2,
    '批判思辨': 3,
    '建立模型': 4,
    '觀察與定題': 5,
    '計畫與執行': 6,
    '分析與發現': 7,
    '討論與傳達': 8,
    '培養科學探究的興趣': 9,
    '養成應用科學思考與探究的習慣': 10,
    '認識科學的本質': 11,
}

# col0 (學習表現項目) label is at the vMerge restart row for each group
# shade col0 if ANY row in the group is selected
GROUP_LABEL_ROW = {
    1: frozenset([1, 2, 3, 4]),   # 探究能力-思考智能
    5: frozenset([5, 6, 7, 8]),   # 探究能力-問題解決
    9: frozenset([9, 10, 11]),    # 科學的態度與本質
}


def highlight_table6_in_doc(doc, selected_subitems: set):
    """
    Shade selected 學習表現子項 rows yellow in an open python-docx Document object.

    Args:
        doc: an open python-docx Document (not yet saved).
        selected_subitems: set of 學習表現子項 strings matching SUBITEM_ROW keys above.
                           Strings NOT in SUBITEM_ROW are silently ignored.
    """
    t = doc.tables[6]
    trs = get_tr_list(t)
    selected_row_indices = {SUBITEM_ROW[s] for s in selected_subitems if s in SUBITEM_ROW}

    # shade col1 (子項) for each selected row
    for ridx in selected_row_indices:
        set_tc_shade(get_tc(trs[ridx], 1), HIGHLIGHT)

    # shade col0 (項目 label, vMerge restart cell) if ANY row in its group is selected
    for label_ridx, group in GROUP_LABEL_ROW.items():
        if group & selected_row_indices:
            set_tc_shade(get_tc(trs[label_ridx], 0), HIGHLIGHT)


def collect_selected_from_table7(doc):
    """
    Convenience helper: read the 子項 column (col index 2) of table7 (學習表現檢核點)
    and return the unique set of 子項 strings found there, to use directly as
    selected_subitems for highlight_table6_in_doc().

    This avoids having to manually track which sub-items were selected — just
    collect them from the table7 content that was already filled in.
    """
    from docx.oxml.ns import qn
    t = doc.tables[7]
    trs = get_tr_list(t)
    selected = set()
    for tr in trs[1:]:  # skip header row
        tc = get_tc(tr, 2)
        texts = tc.findall('.//' + qn('w:t'))
        txt = ''.join(x.text or '' for x in texts).strip()
        if txt and txt in SUBITEM_ROW:
            selected.add(txt)
    return selected
